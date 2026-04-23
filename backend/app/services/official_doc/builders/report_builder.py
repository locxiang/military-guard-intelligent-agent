"""
报告/汇报公文 builder
依据 GB/T 9704-2012 和 GB/T 33476.2-2016
"""
from typing import Dict, Any
from docx import Document

from app.services.official_doc.builders.base import BaseDocumentBuilder
from app.services.official_doc.formats.gb_t_9704_2012 import GB9704_2012


class ReportDocumentBuilder(BaseDocumentBuilder):
    """报告/汇报公文构建器"""

    def build(self, content: Dict[str, Any]) -> Document:
        """
        构建报告公文

        Args:
            content: 内容字典，包含：
                - title: 标题
                - doc_number: 发文字号（可选）
                - recipient: 主送机关
                - situation: 基本情况
                - problems: 存在问题（可选）
                - plan: 下步计划（可选）
                - sender: 发文机关
                - date: 成文日期
                - copies: 抄送机关列表（可选）
                - issuer: 印发机关（可选）
                - issue_date: 印发日期（可选）
                - org_name: 发文机关标志（可选）

        Returns:
            docx Document 对象
        """
        # 红头（发文机关标志）
        self._add_red_header(content.get("org_name", "中国人民解放军XX单位"))

        # 发文字号（可选）
        if content.get("doc_number"):
            self._add_doc_number(content["doc_number"])

        # 红色分隔线
        self._add_red_line()

        # 标题
        if content.get("title"):
            self._add_title(content["title"])

        # 主送机关
        if content.get("recipient"):
            self._add_recipient(content["recipient"])

        # 基本情况
        if content.get("situation"):
            self._add_paragraph(
                "一、基本情况",
                font_name=GB9704_2012.FONT_HEITI,
                font_size=GB9704_2012.FONT_SIZE_MAIN_BODY,
                alignment=GB9704_2012.ALIGN_LEFT,
                first_line_indent=GB9704_2012.FIRST_LINE_INDENT,
                bold=True,
                line_spacing_fixed=GB9704_2012.LINE_SPACING_FIXED,
                zero_spacing=True,
                keep_with_next=True,
            )
            self._add_main_body(content["situation"])

        # 存在问题（可选）
        if content.get("problems"):
            self._add_paragraph(
                "二、存在问题",
                font_name=GB9704_2012.FONT_HEITI,
                font_size=GB9704_2012.FONT_SIZE_MAIN_BODY,
                alignment=GB9704_2012.ALIGN_LEFT,
                first_line_indent=GB9704_2012.FIRST_LINE_INDENT,
                bold=True,
                line_spacing_fixed=GB9704_2012.LINE_SPACING_FIXED,
                zero_spacing=True,
                keep_with_next=True,
            )
            self._add_main_body(content["problems"])

        # 下步计划（可选）
        if content.get("plan"):
            self._add_paragraph(
                "三、下步计划",
                font_name=GB9704_2012.FONT_HEITI,
                font_size=GB9704_2012.FONT_SIZE_MAIN_BODY,
                alignment=GB9704_2012.ALIGN_LEFT,
                first_line_indent=GB9704_2012.FIRST_LINE_INDENT,
                bold=True,
                line_spacing_fixed=GB9704_2012.LINE_SPACING_FIXED,
                zero_spacing=True,
                keep_with_next=True,
            )
            self._add_main_body(content["plan"])

        # 特此报告
        self._add_paragraph(
            "特此报告。",
            font_name=GB9704_2012.FONT_FANGSONG,
            font_size=GB9704_2012.FONT_SIZE_MAIN_BODY,
            alignment=GB9704_2012.ALIGN_LEFT,
            first_line_indent=GB9704_2012.FIRST_LINE_INDENT,
            line_spacing_fixed=GB9704_2012.LINE_SPACING_FIXED,
            zero_spacing=True,
        )

        # 发文机关署名和成文日期
        sender = content.get("sender", "XX单位")
        date = content.get("date", "2026年1月1日")
        self._add_sender_and_date(sender, date)

        # 版记（可选）
        copies = content.get("copies")
        issuer = content.get("issuer")
        issue_date = content.get("issue_date")
        if copies or (issuer and issue_date):
            self._add_version_record(
                copies=copies if isinstance(copies, list) else None,
                issuer=issuer,
                issue_date=issue_date
            )

        # 添加页码
        self._add_page_numbers()

        return self.doc

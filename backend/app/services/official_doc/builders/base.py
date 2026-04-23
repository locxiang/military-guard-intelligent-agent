"""
公文构建器基类
依据 GB/T 33476.2-2016 第2部分：显现
"""
from abc import ABC, abstractmethod
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Cm, Pt, Mm, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

from app.services.official_doc.formats.gb_t_9704_2012 import GB9704_2012


class BaseDocumentBuilder(ABC):
    """公文构建器基类"""

    def __init__(self):
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        """设置页面格式（GB/T 9704-2012 第6章）"""
        section = self.doc.sections[0]

        # 6.1 页面尺寸 - A4 210mm × 297mm
        section.page_width = GB9704_2012.PAGE_WIDTH
        section.page_height = GB9704_2012.PAGE_HEIGHT

        # 6.2 页边距
        section.top_margin = GB9704_2012.MARGIN_TOP
        section.bottom_margin = GB9704_2012.MARGIN_BOTTOM
        section.left_margin = GB9704_2012.MARGIN_LEFT
        section.right_margin = GB9704_2012.MARGIN_RIGHT

        # 设置页眉页脚距离
        section.header_distance = Mm(15)
        section.footer_distance = GB9704_2012.PAGE_NUMBER_CENTER_DISTANCE

    def _get_chinese_space(self, num_chars: int) -> str:
        """
        获取指定数量的汉字空格（全角空格）

        Args:
            num_chars: 汉字数量

        Returns:
            空格字符串
        """
        return '\u3000' * num_chars

    def _set_font(self, run, font_name: str, font_size: Pt, bold: bool = False, color: RGBColor = None):
        """
        设置字体（支持中文字体）

        Args:
            run: docx Run 对象
            font_name: 字体名称
            font_size: 字号
            bold: 是否加粗
            color: 字体颜色
        """
        run.font.name = font_name
        run.font.size = font_size
        run.font.bold = bold
        if color:
            run.font.color.rgb = color
        # 设置中文字体
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _add_paragraph(
        self,
        text: str = "",
        font_name: str = GB9704_2012.FONT_FANGSONG,
        font_size: Pt = GB9704_2012.FONT_SIZE_MAIN_BODY,
        alignment=GB9704_2012.ALIGN_LEFT,
        first_line_indent: Mm = None,
        left_indent: Mm = None,
        right_indent: Mm = None,
        space_before: Mm = None,
        space_before_twips: int = None,
        space_after: Mm = None,
        space_after_twips: int = None,
        line_spacing: float = None,
        line_spacing_fixed: Pt = None,
        bold: bool = False,
        color: RGBColor = None,
        keep_with_next: bool = False,
        zero_spacing: bool = False,
    ):
        """
        添加段落（精确间距控制）

        Args:
            text: 文本内容
            font_name: 字体名称
            font_size: 字号
            alignment: 对齐方式
            first_line_indent: 首行缩进 (Mm)
            left_indent: 左缩进 (Mm)
            right_indent: 右缩进 (Mm)
            space_before: 段前间距 (Mm)
            space_before_twips: 段前间距 (twips，优先使用)
            space_after: 段后间距 (Mm)
            space_after_twips: 段后间距 (twips，优先使用)
            line_spacing: 行距倍数
            line_spacing_fixed: 固定行距 (Pt)
            bold: 是否加粗
            color: 字体颜色
            keep_with_next: 是否与下段同页
            zero_spacing: 是否段前段后都为0

        Returns:
            docx Paragraph 对象
        """
        from docx.enum.text import WD_LINE_SPACING

        p = self.doc.add_paragraph()
        p.alignment = alignment

        # 缩进设置
        if first_line_indent:
            p.paragraph_format.first_line_indent = first_line_indent
        if left_indent:
            p.paragraph_format.left_indent = left_indent
        if right_indent:
            p.paragraph_format.right_indent = right_indent

        # 间距设置
        if zero_spacing:
            p.paragraph_format.space_before = 0
            p.paragraph_format.space_after = 0
        else:
            if space_before_twips is not None:
                p.paragraph_format.space_before = space_before_twips
            elif space_before:
                p.paragraph_format.space_before = GB9704_2012.mm_to_twips(space_before.mm)

            if space_after_twips is not None:
                p.paragraph_format.space_after = space_after_twips
            elif space_after:
                p.paragraph_format.space_after = GB9704_2012.mm_to_twips(space_after.mm)

        # 行距设置
        if line_spacing_fixed:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = line_spacing_fixed
        elif line_spacing:
            p.paragraph_format.line_spacing = line_spacing

        # 分页控制
        if keep_with_next:
            p.paragraph_format.keep_with_next = True

        # 添加文本
        if text:
            run = p.add_run(text)
            self._set_font(run, font_name, font_size, bold=bold, color=color)

        return p

    def _add_red_header(self, org_name: str = "中国人民解放军XX单位"):
        """
        添加红头（发文机关标志）
        7.1 版头：上边缘至版心上边缘 37.42mm

        Args:
            org_name: 发文机关名称
        """
        # 计算需要的段前间距
        space_before_mm = GB9704_2012.RED_HEADER_TOP_MARGIN.mm - GB9704_2012.MARGIN_TOP.mm
        space_before_twips = GB9704_2012.mm_to_twips(space_before_mm) if space_before_mm > 0 else 0

        # 红头文字（发文机关标志使用方正小标宋简体）
        p = self._add_paragraph(
            text=org_name,
            font_name=GB9704_2012.FONT_XIAOBIAOSONG,
            font_size=GB9704_2012.FONT_SIZE_RED_HEADER,
            alignment=GB9704_2012.ALIGN_CENTER,
            bold=True,
            color=GB9704_2012.COLOR_RED,
            space_before_twips=space_before_twips,
            keep_with_next=True,
        )

    def _add_doc_number(self, doc_number: str):
        """
        添加发文字号
        7.1 版头：上边缘至发文机关名称下边缘 2行 (20.78mm)

        Args:
            doc_number: 发文字号，如：〔2026〕保字第 001 号
        """
        self._add_paragraph(
            text=doc_number,
            font_name=GB9704_2012.FONT_FANGSONG,
            font_size=GB9704_2012.FONT_SIZE_DOC_NUMBER,
            alignment=GB9704_2012.ALIGN_RIGHT,
            space_before=GB9704_2012.DOC_NUMBER_SPACING_AFTER_HEADER,
            keep_with_next=True,
        )

    def _add_red_line(self):
        """添加红色分隔线（GB/T 9704-2012）
        7.1 版头：红色分隔线上边缘至发文字号下边缘 4mm
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = GB9704_2012.mm_to_twips(4)
        p.paragraph_format.keep_with_next = True

        # 设置段落底部边框为红色粗实线
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 线宽，1/8 磅，12 = 1.5 磅
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'FF0000')
        pBdr.append(bottom)
        pPr.append(pBdr)

        # 添加空内容
        run = p.add_run()
        run.font.size = Pt(1)

    def _add_title(self, title: str):
        """
        添加公文标题
        7.2 主体：上边缘至红色分隔线 2行 (20.78mm)

        Args:
            title: 标题文本
        """
        self._add_paragraph(
            text=title,
            font_name=GB9704_2012.FONT_XIAOBIAOSONG,
            font_size=GB9704_2012.FONT_SIZE_TITLE,
            alignment=GB9704_2012.ALIGN_CENTER,
            bold=True,
            space_before=GB9704_2012.TITLE_SPACING_AFTER_RED_LINE,
            keep_with_next=True,
        )

    def _add_recipient(self, recipient: str):
        """
        添加主送机关
        7.2 主体：上边缘至标题下边缘 1行 (10.39mm)

        Args:
            recipient: 主送机关名称
        """
        if not recipient.endswith("：") and not recipient.endswith(":"):
            recipient = recipient + "："
        self._add_paragraph(
            text=recipient,
            font_name=GB9704_2012.FONT_KAITI,
            font_size=GB9704_2012.FONT_SIZE_RECIPIENT,
            alignment=GB9704_2012.ALIGN_LEFT,
            space_before=GB9704_2012.RECIPIENT_SPACING_AFTER_TITLE,
            keep_with_next=True,
        )

    def _add_main_body(self, content: str):
        """
        添加正文
        7.2 主体：首行缩进 2 个三号字，层级编号：一、→（一）→1.→（1）

        Args:
            content: 正文内容（按换行符分段）
        """
        import re
        from docx.enum.text import WD_LINE_SPACING

        paragraphs = content.split("\n")
        for i, para_text in enumerate(paragraphs):
            if para_text.strip():
                is_last = (i == len(paragraphs) - 1)

                # 检查是否包含分段标题，匹配到第一个句号、冒号、分号
                title_end_pos = -1
                # 匹配模式：一、 二、 三、 ... 十、
                patterns = [
                    r'^[一二三四五六七八九十]+、',
                    r'^\d+[.、]',
                    r'^（[一二三四五六七八九十]+）',
                    r'^\([一二三四五六七八九十]+\)',
                    r'^\d+）'
                ]

                has_title = False
                for pattern in patterns:
                    if re.match(pattern, para_text.strip()):
                        has_title = True
                        break

                if has_title:
                    # 找到第一个句号、冒号、分号的位置作为标题结束
                    for idx, char in enumerate(para_text):
                        if char in ['。', '：', '；', ':', ';']:
                            title_end_pos = idx + 1
                            break

                    # 如果没找到标点，把整个段落都作为标题
                    if title_end_pos == -1:
                        title_end_pos = len(para_text)

                    # 分离标题和正文
                    title_part = para_text[:title_end_pos]
                    body_part = para_text[title_end_pos:]

                    # 检查标题长度，超过20个字就不处理成黑体
                    if len(title_part.strip()) > 20:
                        # 超过20个字，当作普通正文处理
                        self._add_paragraph(
                            text=para_text,
                            font_name=GB9704_2012.FONT_FANGSONG,
                            font_size=GB9704_2012.FONT_SIZE_MAIN_BODY,
                            alignment=GB9704_2012.ALIGN_JUSTIFY,
                            first_line_indent=GB9704_2012.FIRST_LINE_INDENT,
                            line_spacing_fixed=GB9704_2012.LINE_SPACING_FIXED,
                            zero_spacing=True,
                            keep_with_next=not is_last,
                        )
                    else:
                        # 不超过20个字，处理成黑体标题
                        # 创建段落，分别添加标题和正文
                        p = self.doc.add_paragraph()
                        p.alignment = GB9704_2012.ALIGN_LEFT
                        p.paragraph_format.first_line_indent = GB9704_2012.FIRST_LINE_INDENT
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                        p.paragraph_format.line_spacing = GB9704_2012.LINE_SPACING_FIXED
                        p.paragraph_format.space_before = 0
                        p.paragraph_format.space_after = 0
                        if not is_last:
                            p.paragraph_format.keep_with_next = True

                        # 标题部分使用黑体（从开头到第一个标点符号）
                        if title_part:
                            run = p.add_run(title_part)
                            self._set_font(run, GB9704_2012.FONT_HEITI, GB9704_2012.FONT_SIZE_MAIN_BODY, bold=True)

                        # 正文部分使用仿宋
                        if body_part.strip():
                            run = p.add_run(body_part)
                            self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_MAIN_BODY)
                else:
                    # 普通正文使用仿宋，段前、段后为0，固定行距28.95磅，首行缩进2字符
                    self._add_paragraph(
                        text=para_text,
                        font_name=GB9704_2012.FONT_FANGSONG,
                        font_size=GB9704_2012.FONT_SIZE_MAIN_BODY,
                        alignment=GB9704_2012.ALIGN_JUSTIFY,
                        first_line_indent=GB9704_2012.FIRST_LINE_INDENT,
                        line_spacing_fixed=GB9704_2012.LINE_SPACING_FIXED,
                        zero_spacing=True,
                        keep_with_next=not is_last,
                    )

    def _add_attachment(self, attachment: str):
        """
        添加附件说明
        7.2 主体：上边缘至正文末段下边缘 1行 (10.39mm)；回行首字对齐

        Args:
            attachment: 附件说明
        """
        from docx.enum.text import WD_LINE_SPACING

        p = self._add_paragraph(
            font_name=GB9704_2012.FONT_FANGSONG,
            font_size=GB9704_2012.FONT_SIZE_MAIN_BODY,
            alignment=GB9704_2012.ALIGN_LEFT,
            space_before=GB9704_2012.ATTACHMENT_SPACING_AFTER_BODY,
        )

        # 应用段前、段后为0，固定行距28.95磅
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = GB9704_2012.LINE_SPACING_FIXED

        # "附件："
        run = p.add_run("附件：")
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_MAIN_BODY)

        # 附件内容
        run = p.add_run(attachment)
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_MAIN_BODY)

    def _add_sender_and_date(self, sender: str, date: str):
        """
        添加发文机关署名和成文日期（右空四字，GB/T 9704-2012）
        7.2 主体：成文日期右缩进 4 字

        Args:
            sender: 发文机关
            date: 成文日期
        """
        from docx.enum.text import WD_LINE_SPACING

        # 计算右缩进：4个三号字宽
        right_indent = Mm(5.54 * 4)  # 每个三号字约 5.54mm

        # 发文机关署名
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_indent = right_indent
        p.paragraph_format.space_before = GB9704_2012.mm_to_twips(20.78)  # 2行
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = GB9704_2012.LINE_SPACING_FIXED
        p.paragraph_format.keep_with_next = True

        run = p.add_run(sender)
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_MAIN_BODY)

        # 成文日期
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.right_indent = right_indent
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = GB9704_2012.LINE_SPACING_FIXED

        run = p.add_run(date)
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_MAIN_BODY)

    def _add_version_separator_line(self, is_first: bool = False, is_last: bool = False):
        """添加版记分隔线（细实线，GB/T 9704-2012）"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 设置段落边框
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        if is_first or not is_last:
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), '4')  # 细线，0.5 磅
            top.set(qn('w:space'), '0')
            top.set(qn('w:color'), '000000')
            pBdr.append(top)

        if is_last or not is_first:
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '4')
            bottom.set(qn('w:space'), '0')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)

        pPr.append(pBdr)

        # 添加空内容
        run = p.add_run()
        run.font.size = Pt(1)

    def _add_copy_recipients(self, copies: List[str]):
        """
        添加抄送机关（左空一字，回行对齐，GB/T 9704-2012）
        7.3 版记："抄送："左缩进 1 字；回行首字对齐

        Args:
            copies: 抄送机关列表
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 左缩进 1 字
        left_indent = GB9704_2012.COPY_RECIPIENT_LEFT_INDENT
        p.paragraph_format.left_indent = left_indent

        # "抄送："
        run = p.add_run("抄送：")
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_FOOTER)

        # 抄送机关列表
        copy_text = "，".join(copies) + "。"
        run = p.add_run(copy_text)
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_FOOTER)

    def _add_issuer_and_date(self, issuer: str, issue_date: str):
        """
        添加印发机关和日期（GB/T 9704-2012）
        7.3 版记：印发机关左缩进 1 字，印发日期右缩进 1 字

        Args:
            issuer: 印发机关
            issue_date: 印发日期
        """
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # 使用制表位实现左右对齐
        section = self.doc.sections[-1]
        page_width = section.page_width
        left_margin = section.left_margin
        right_margin = section.right_margin
        effective_width = page_width - left_margin - right_margin

        # 添加制表位
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(effective_width, WD_ALIGN_PARAGRAPH.RIGHT)

        # 左空一字 - 印发机关
        left_indent = GB9704_2012.ISSUER_LEFT_INDENT
        p.paragraph_format.left_indent = left_indent

        run = p.add_run(issuer)
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_FOOTER)

        # 制表符
        run = p.add_run("\t")
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_FOOTER)

        # 印发日期
        run = p.add_run(issue_date)
        self._set_font(run, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_FOOTER)

    def _add_version_record(self, copies: List[str] = None, issuer: str = None, issue_date: str = None):
        """
        添加完整的版记部分（GB/T 9704-2012）
        7.3 版记：排在偶数页，末条分隔线与版心下边缘重合

        Args:
            copies: 抄送机关列表（可选）
            issuer: 印发机关（可选）
            issue_date: 印发日期（可选）
        """
        # 确保版记在偶数页
        if len(self.doc.paragraphs) > 0:
            last_para = self.doc.paragraphs[-1]
            last_para.paragraph_format.page_break_after = True

        # 版记上分隔线
        self._add_version_separator_line(is_first=True)

        # 抄送机关（如果有）
        if copies:
            self._add_copy_recipients(copies)
            # 版记中分隔线
            self._add_version_separator_line()

        # 印发机关和日期
        if issuer and issue_date:
            self._add_issuer_and_date(issuer, issue_date)

        # 版记下分隔线（末条）
        self._add_version_separator_line(is_last=True)

    def _add_page_numbers(self):
        """
        添加页码（GB/T 9704-2012 第8章）
        - 版心下边缘至页码中心：4.58mm
        - 单页码：右对齐、右缩进 1 字
        - 双页码：左对齐、左缩进 1 字
        """
        # 为文档添加不同的页眉页脚（奇数页和偶数页）
        section = self.doc.sections[0]
        section.different_odd_and_even_pages_header_footer = True

        # 奇数页页码（右对齐）
        footer_odd = section.footer
        p_odd = footer_odd.paragraphs[0] if footer_odd.paragraphs else footer_odd.add_paragraph()
        p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_odd.paragraph_format.right_indent = GB9704_2012.PAGE_NUMBER_SIDE_INDENT

        # 添加页码字段
        run_odd = p_odd.add_run()
        self._set_font(run_odd, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_FOOTER)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_odd._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run_odd._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_odd._r.append(fldChar2)

        run_odd.add_text("1")

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_odd._r.append(fldChar3)

        # 偶数页页码（左对齐）
        footer_even = section.even_page_footer
        p_even = footer_even.paragraphs[0] if footer_even.paragraphs else footer_even.add_paragraph()
        p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_even.paragraph_format.left_indent = GB9704_2012.PAGE_NUMBER_SIDE_INDENT

        # 添加页码字段
        run_even = p_even.add_run()
        self._set_font(run_even, GB9704_2012.FONT_FANGSONG, GB9704_2012.FONT_SIZE_FOOTER)
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run_even._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        run_even._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        run_even._r.append(fldChar2)

        run_even.add_text("1")

        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        run_even._r.append(fldChar3)

    @abstractmethod
    def build(self, content: Dict[str, Any]) -> Document:
        """
        构建公文（子类实现）

        Args:
            content: 内容字典

        Returns:
            docx Document 对象
        """
        pass

    def save(self, file_path: str):
        """
        保存文档

        Args:
            file_path: 文件路径
        """
        self.doc.save(file_path)

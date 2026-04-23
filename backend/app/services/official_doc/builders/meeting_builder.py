"""
会议纪要 builder
"""
from typing import Dict, Any
from docx import Document

from app.services.official_doc.builders.base import BaseDocumentBuilder


class MeetingMinutesBuilder(BaseDocumentBuilder):
    """会议纪要构建器"""

    def build(self, content: Dict[str, Any]) -> Document:
        """
        构建会议纪要

        Args:
            content: 内容字典

        Returns:
            docx Document 对象
        """
        # 标题
        if content.get("title"):
            self._add_title(content["title"])

        # 会议基本信息
        if content.get("meeting_info"):
            self._add_paragraph("一、会议基本信息")
            self._add_main_body(content["meeting_info"])

        # 会议议题
        if content.get("topics"):
            self._add_paragraph("二、会议议题")
            self._add_main_body(content["topics"])

        # 讨论内容
        if content.get("discussion"):
            self._add_paragraph("三、讨论内容")
            self._add_main_body(content["discussion"])

        # 议定事项
        if content.get("decisions"):
            self._add_paragraph("四、议定事项")
            self._add_main_body(content["decisions"])

        # 待办分工（可选）
        if content.get("tasks"):
            self._add_paragraph("五、待办分工")
            self._add_main_body(content["tasks"])

        return self.doc
